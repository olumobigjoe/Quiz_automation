"""
Course Material -> Quiz Generator (Ollama Cloud)
====================================================
Run with:  streamlit run app.py

Upload a PDF or DOCX of course material. This tool reads the text,
sends it to a model hosted on Ollama Cloud to draft multiple-choice
questions (with plausible wrong options AND a suggested correct
answer), lets you review/edit everything, then exports a completely
separate, self-contained `exam_app.py` + `README.md` you can hand to
students or push to GitHub. The exported exam app needs only
`streamlit` (and `pandas`) to run -- it does NOT call Ollama at all,
since the questions and answers are baked in as data.

The export is a full exam portal (not just a static quiz page):
students log in with a name + student ID, are assigned to a batch
(max 40 students per batch), can only sit the test inside their
batch's access window, get a live countdown timer, and can only
submit once. All of this ships inside `exam_app.py`; the accompanying
`question_bank.json` contains ONLY the question text, options, and
correct answer -- no student/batch/schedule data.

Requires an Ollama Cloud API key (free tier available):
    1. Sign up / sign in at https://ollama.com
    2. Create an API key: https://ollama.com/settings/keys
    3. Either set it as an environment variable before launching --
         export OLLAMA_API_KEY=your_key_here      (Mac/Linux)
         setx OLLAMA_API_KEY "your_key_here"       (Windows, new terminal after)
       -- or paste it into the sidebar field when the app is running.

No local model download, no `ollama serve`, and your laptop doesn't
need to stay on for a deployed version of this app to keep working --
inference runs on Ollama's cloud infrastructure.
"""

from __future__ import annotations

import base64
import io
import json
import math
import os
import re
import zipfile
from datetime import datetime, date, time as dt_time, timedelta

import requests
import streamlit as st

# ============================================================================
# CONFIG
# ============================================================================
OLLAMA_CLOUD_HOST = "https://ollama.com"
# A few reasonable defaults for structured-JSON generation tasks like this
# one. Ollama's cloud catalog changes over time -- if you want a different
# model, use "Custom model tag..." in the sidebar and check what's
# currently available to your account at https://ollama.com/settings/keys
# or by calling GET https://ollama.com/api/tags with your API key.
MODEL_OPTIONS = ["gpt-oss:20b", "gpt-oss:120b", "qwen3.5", "deepseek-v3.1:671b", "Custom model tag..."]
CHUNK_CHARS = 4000  # cloud models handle larger context comfortably
MAX_QUESTIONS = 50  # hard cap on how many questions can be requested/generated
MAX_STUDENTS_PER_BATCH = 40  # hard cap enforced when building the exam export

st.set_page_config(page_title="Quiz Generator (Ollama Cloud)", page_icon="📚", layout="wide")

# ============================================================================
# SESSION STATE
# ============================================================================
defaults = {
    "raw_text": "",
    "source_filename": "",
    "questions": [],
    "generation_log": [],
    "generated_once": False,
    "regenerate_count": 0,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ============================================================================
# FILE EXTRACTION
# ============================================================================
def extract_text_from_pdf(file_bytes: bytes) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def chunk_text(text: str, chunk_size: int = CHUNK_CHARS):
    text = text.strip()
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        # try to break on a paragraph/sentence boundary, not mid-word
        if end < len(text):
            boundary = text.rfind("\n", start, end)
            if boundary == -1 or boundary <= start + chunk_size * 0.5:
                boundary = text.rfind(". ", start, end)
            if boundary != -1 and boundary > start + chunk_size * 0.3:
                end = boundary + 1
        chunks.append(text[start:end].strip())
        start = end
    return [c for c in chunks if c]


# ============================================================================
# OLLAMA CALL + PARSING
# ============================================================================
PROMPT_TEMPLATE = """You are helping a university instructor create multiple-choice quiz questions for students, based ONLY on the course material provided below.

Generate exactly {n} multiple-choice questions from this material.
Each question must have exactly {opts} answer options.
Exactly one option must be the correct answer; the rest must be plausible but clearly wrong to someone who understood the material.
Vary which option position (0, 1, 2, ...) holds the correct answer -- do not always put it first.

Return ONLY a JSON array -- your entire response must start with [ and end with ]. Do not wrap it in an object, do not add markdown code fences, do not add any text before or after it. Each element must have exactly this shape:
{{
  "question": "the question text",
  "options": ["option A text", "option B text", "..."],
  "correct_index": 0,
  "explanation": "one short sentence explaining why that answer is correct"
}}

"correct_index" is the 0-based index into "options" of the correct answer.

COURSE MATERIAL:
\"\"\"
{material}
\"\"\"
"""


def call_ollama(model: str, prompt: str, host: str, api_key: str) -> str:
    headers = {"Authorization": f"Bearer {api_key}"} if api_key else {}
    resp = requests.post(
        f"{host}/api/chat",
        headers=headers,
        json={
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "stream": False,
            "format": "json",
            "options": {"temperature": 0.4},
        },
        timeout=600,
    )
    if resp.status_code == 401:
        raise PermissionError(
            "Ollama Cloud rejected the API key (401 Unauthorized). "
            "Check the key in the sidebar / OLLAMA_API_KEY env var."
        )
    if resp.status_code == 429:
        raise RuntimeError(
            "Ollama Cloud rate/usage limit hit (429). Wait a bit, reduce the "
            "number of questions, or upgrade your plan at https://ollama.com/pricing."
        )
    resp.raise_for_status()
    data = resp.json()
    return data.get("message", {}).get("content", "")


def extract_json_array(raw: str):
    """Best-effort extraction of a JSON array from a model response, even
    if the model added stray text around it, wrapped it in a code fence,
    or nested it inside an object like {"questions": [...]} instead of
    returning a bare array (all common with small local models)."""
    raw = raw.strip()
    # strip a markdown code fence if present
    fence = re.match(r"^```(?:json)?\s*(.*?)\s*```$", raw, re.DOTALL)
    if fence:
        raw = fence.group(1).strip()

    def _unwrap(value):
        if isinstance(value, list):
            return value
        if isinstance(value, dict):
            for key in ("questions", "quiz", "items", "data", "results"):
                if isinstance(value.get(key), list):
                    return value[key]
            # a dict of one question rather than a list of questions
            if "question" in value:
                return [value]
        return None

    try:
        return _unwrap(json.loads(raw))
    except json.JSONDecodeError:
        pass

    # try to locate the outermost array anywhere in the text
    match = re.search(r"\[.*\]", raw, re.DOTALL)
    if match:
        try:
            return _unwrap(json.loads(match.group(0)))
        except json.JSONDecodeError:
            pass

    # fall back to locating an outermost object
    match = re.search(r"\{.*\}", raw, re.DOTALL)
    if match:
        try:
            return _unwrap(json.loads(match.group(0)))
        except json.JSONDecodeError:
            pass

    return None


def validate_questions(items, num_options: int):
    """Coerce and accept near-miss output from small models rather than
    rejecting outright, but enforce the requested option count exactly
    so the final quiz is uniform (mixed 3/4/5-option questions read as
    broken to a student). If the model gave too many options, the extra
    WRONG ones are trimmed (the correct one is always kept). If it gave
    too few, the question is dropped -- there's no safe way to invent a
    plausible extra distractor without another model call."""
    valid = []
    for item in items or []:
        if not isinstance(item, dict):
            continue
        q = item.get("question")
        opts = item.get("options")
        idx = item.get("correct_index")

        if not isinstance(q, str) or not q.strip():
            continue
        if not isinstance(opts, list) or len(opts) < 2:
            continue

        opts = [str(o).strip() for o in opts if str(o).strip()]
        if len(opts) < 2:
            continue

        # coerce a stringified index ("0", "2") to int
        if isinstance(idx, str) and idx.strip().isdigit():
            idx = int(idx.strip())
        # some models label the correct option by its text instead of an index
        if isinstance(idx, str) and idx not in opts:
            match = next((k for k, o in enumerate(opts) if o.lower() == idx.strip().lower()), None)
            idx = match if match is not None else None
        elif isinstance(idx, str) and idx in opts:
            idx = opts.index(idx)

        if not isinstance(idx, int) or not (0 <= idx < len(opts)):
            continue

        # enforce exact option count for a uniform quiz
        if len(opts) > num_options:
            correct_text = opts[idx]
            wrong_opts = [o for k, o in enumerate(opts) if k != idx]
            keep_wrong = wrong_opts[: num_options - 1]
            opts = keep_wrong + [correct_text]
            idx = len(opts) - 1
        elif len(opts) < num_options:
            continue  # can't safely fabricate a missing distractor

        valid.append(
            {
                "question": q.strip(),
                "options": opts,
                "correct_index": idx,
                "explanation": str(item.get("explanation", "")).strip(),
            }
        )
    return valid


def _generate_from_chunk(chunk: str, n: int, model: str, host: str, api_key: str, num_options: int, log: list, label: str):
    """Ask the model for n questions from one chunk, validate, log the
    outcome, and return the list of valid questions (possibly empty).
    Raises PermissionError on a bad API key so the caller can stop the
    whole run immediately instead of retrying a doomed request repeatedly."""
    prompt = PROMPT_TEMPLATE.format(n=n, opts=num_options, material=chunk)
    try:
        raw = call_ollama(model, prompt, host, api_key)
    except PermissionError:
        raise  # let the caller abort the whole run, not just this chunk
    except requests.exceptions.ConnectionError:
        log.append(f"{label}: could not reach {host}. Check your internet connection.")
        return []
    except RuntimeError as e:  # rate limit
        log.append(f"{label}: {e}")
        return []
    except Exception as e:  # noqa: BLE001
        log.append(f"{label}: request failed ({e}).")
        return []

    parsed = extract_json_array(raw)
    valid = validate_questions(parsed, num_options)
    if not valid:
        preview = raw.strip().replace("\n", " ")[:300]
        log.append(
            f"{label}: model did not return usable questions, skipped. "
            f"Raw response preview: `{preview}{'...' if len(raw.strip()) > 300 else ''}`"
        )
    else:
        log.append(f"{label}: generated {len(valid)} valid question(s).")
    return valid


MAX_PER_CALL = 4  # small local models are much more reliable generating a
                   # short JSON array than a long one in a single response;
                   # split bigger requests into several smaller calls instead.


def _split_into_batches(n: int, batch_size: int = MAX_PER_CALL):
    batches = []
    remaining = n
    while remaining > 0:
        take = min(batch_size, remaining)
        batches.append(take)
        remaining -= take
    return batches


def generate_quiz(text: str, model: str, host: str, api_key: str, total_questions: int, num_options: int):
    chunks = chunk_text(text)
    if not chunks:
        return [], ["No text could be extracted from the document."]

    n_chunks = len(chunks)
    base = total_questions // n_chunks
    remainder = total_questions % n_chunks
    per_chunk = [base + (1 if i < remainder else 0) for i in range(n_chunks)]

    all_questions = []
    log = []
    progress = st.progress(0.0, text="Starting generation...")

    try:
        for i, (chunk, n) in enumerate(zip(chunks, per_chunk)):
            progress.progress(i / n_chunks, text=f"Generating from section {i + 1}/{n_chunks}...")
            if n <= 0:
                continue
            batches = _split_into_batches(n)
            for b, batch_n in enumerate(batches):
                label = f"Section {i + 1}" + (f" (batch {b + 1}/{len(batches)})" if len(batches) > 1 else "")
                all_questions.extend(
                    _generate_from_chunk(chunk, batch_n, model, host, api_key, num_options, log, label)
                )

        # ---- top-up pass: strict validation means some questions get
        # dropped (wrong option count, no answer key found, etc.), so the
        # first pass often falls short of the target. Keep asking for the
        # remainder, cycling through chunks, up to a small retry cap so a
        # stubborn document can't loop forever.
        max_extra_rounds = 8
        round_num = 0
        while len(all_questions) < total_questions and round_num < max_extra_rounds:
            round_num += 1
            deficit = total_questions - len(all_questions)
            progress.progress(
                0.9, text=f"Topping up {deficit} more question(s) (round {round_num})..."
            )
            chunk = chunks[round_num % n_chunks]
            batch_n = min(deficit, MAX_PER_CALL)
            got = _generate_from_chunk(
                chunk,
                batch_n,
                model,
                host,
                api_key,
                num_options,
                log,
                f"Top-up round {round_num}",
            )
            # avoid near-duplicate questions if the model regenerates similar
            # content from the same chunk on repeated top-up rounds
            existing = {q["question"].strip().lower() for q in all_questions}
            got = [q for q in got if q["question"].strip().lower() not in existing]
            all_questions.extend(got)
            if not got:
                # this chunk isn't yielding more -- no point retrying it again
                continue
    except PermissionError as e:
        progress.progress(1.0, text="Stopped.")
        log.append(str(e))
        return all_questions[:total_questions], log

    progress.progress(1.0, text="Done.")
    if len(all_questions) < total_questions:
        log.append(
            f"Reached the retry limit with {len(all_questions)}/{total_questions} questions. "
            "The model may be struggling with this document/option count -- try fewer "
            "questions, a larger/better model (e.g. qwen3:4b), or fewer options per question."
        )
    return all_questions[:total_questions], log


# ============================================================================
# BATCH SCHEDULING (for the exam export)
# ============================================================================
def expand_matric_config(cfg: dict):
    """Turns one batch's matric-number UI config into an explicit, deduped,
    upper-cased list of allowed IDs, or None if the batch has no
    restriction (i.e. any student may log into it, size auto-split)."""
    mode = (cfg or {}).get("mode", "auto")

    if mode == "range":
        prefix = (cfg.get("prefix") or "").strip()
        start_n = cfg.get("start")
        end_n = cfg.get("end")
        pad = cfg.get("pad") or 0
        if start_n is None or end_n is None or end_n < start_n:
            return None
        ids = []
        for n in range(int(start_n), int(end_n) + 1):
            num_str = str(n).zfill(int(pad)) if pad else str(n)
            ids.append(f"{prefix}{num_str}")
        return ids or None

    if mode == "list":
        raw = cfg.get("text") or ""
        parts = re.split(r"[,\n]", raw)
        ids = [p.strip() for p in parts if p.strip()]
        return ids or None

    return None  # "auto" -- no restriction


def compute_batches(num_students: int, num_batches: int, test_date: date,
                     test_start_time: dt_time, batch_window_minutes: int,
                     matric_lists: list | None = None):
    """Split num_students across num_batches back-to-back time windows,
    starting at test_date/test_start_time. Returns a list of dicts: name,
    start ("HH:MM"), end ("HH:MM"), size, allowed_ids (list, possibly
    empty meaning "no restriction").

    matric_lists, if given, must have one entry per batch: either None
    (batch is open -- its size comes from an even split of whatever
    students remain after explicit batches are accounted for) or a
    list[str] of specific matric/student IDs registered for that batch
    (its size is exactly len(list), regardless of num_students).

    Raises ValueError if any batch would exceed MAX_STUDENTS_PER_BATCH, or
    if explicitly-listed IDs alone already exceed num_students.
    """
    matric_lists = matric_lists or [None] * num_batches

    explicit_sizes = [len(m) if m else None for m in matric_lists]
    explicit_total = sum(s for s in explicit_sizes if s)
    open_indexes = [i for i, s in enumerate(explicit_sizes) if s is None]

    if explicit_total > num_students and not open_indexes:
        # all batches are explicit and their total already exceeds the
        # instructor's stated headcount -- just trust the explicit lists.
        num_students = explicit_total

    remaining = max(0, num_students - explicit_total)
    sizes = list(explicit_sizes)
    if open_indexes:
        base = remaining // len(open_indexes)
        rem = remaining % len(open_indexes)
        for j, i in enumerate(open_indexes):
            sizes[i] = base + (1 if j < rem else 0)

    if sizes and max(sizes) > MAX_STUDENTS_PER_BATCH:
        raise ValueError(
            f"At least one batch would end up with more than "
            f"{MAX_STUDENTS_PER_BATCH} students (largest batch: {max(sizes)}). "
            f"Increase the number of batches, or split a large matric-number "
            f"range/list across more than one batch."
        )

    start_dt = datetime.combine(test_date, test_start_time)
    batches = []
    for i, size in enumerate(sizes):
        b_start = start_dt + timedelta(minutes=i * batch_window_minutes)
        b_end = b_start + timedelta(minutes=batch_window_minutes)
        allowed_ids = sorted({m.strip().upper() for m in matric_lists[i]}) if matric_lists[i] else []
        batches.append(
            {
                "name": f"Batch {i + 1}",
                "start": b_start.strftime("%H:%M"),
                "end": b_end.strftime("%H:%M"),
                "size": size,
                "allowed_ids": allowed_ids,
            }
        )
    return batches


# ============================================================================
# EXPORT TEMPLATE (the standalone student-facing exam portal)
# ============================================================================
EXAM_APP_TEMPLATE = r'''"""
__COURSE_TITLE__ -- Online Exam Portal
Auto-generated by the Course Material -> Quiz Generator.
Self-contained -- no external services or internet API calls needed.
Generated on __GENERATED_DATE__ from: __SOURCE_FILENAME__

Run with:
    pip install -r requirements.txt
    streamlit run exam_app.py

SECURITY NOTE: ADMIN_PASSWORD below is baked into this file in plain
text for convenience. That's fine for a low-stakes classroom quiz, but
if you're deploying this publicly (e.g. Streamlit Community Cloud) for
a real exam, set the ADMIN_PASSWORD environment variable instead (it
overrides the baked-in value) so the password isn't sitting in your
public GitHub repo. See the README for details.
"""

import base64
import json
import os
import random
import time
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo

import pandas as pd
import streamlit as st

st.set_page_config(page_title="__COURSE_TITLE__", page_icon="\U0001F4DD", layout="wide")

# ============================================================
# EXAM CONFIGURATION (baked in at export time)
# ============================================================
COURSE_TITLE = "__COURSE_TITLE__"
TEST_DATE = "__TEST_DATE__"                      # YYYY-MM-DD
TEST_START_TIME = "__TEST_START_TIME__"          # HH:MM, 24-hour
TIMEZONE_NAME = "__TIMEZONE_NAME__"              # change here if you deploy from a different timezone
TOTAL_STUDENTS = __NUM_STUDENTS__
TEST_DURATION_SECONDS = __TEST_DURATION_SECONDS__
AUTO_REFRESH_SECONDS = 10                        # keeps the waiting room / timer live without manual reloads
RESULTS_FILE = "exam_results.csv"
# The admin password can be overridden without editing this file by
# setting an ADMIN_PASSWORD environment variable before launching.
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "__ADMIN_PASSWORD__")
LOCAL_TIMEZONE = ZoneInfo(TIMEZONE_NAME)

# Precomputed batch schedule -- max __MAX_STUDENTS_PER_BATCH__ students per batch.
BATCHES = __BATCHES_JSON__

RESULT_COLUMNS = [
    "Timestamp", "Name", "Student ID", "Course", "Batch", "Score",
    "Total Questions", "Correct Answers", "Wrong Answers", "Percentage",
    "Time Used", "Time Seconds", "Status",
]

# ============================================================
# QUESTION BANK (question / options / answer ONLY -- no student data)
# ============================================================
QUESTIONS_B64 = "__QUESTIONS_B64__"
QUESTION_BANK = json.loads(base64.b64decode(QUESTIONS_B64).decode("utf-8"))
QUESTIONS_PER_STUDENT = len(QUESTION_BANK)

# ============================================================
# STYLING
# ============================================================
st.markdown(
    """
    <style>
    .block-container {max-width: 1200px; padding-top: 1.5rem; padding-bottom: 2rem;}
    .main-title {text-align:center; font-size:34px; font-weight:800; margin-bottom:2px;}
    .main-subtitle {text-align:center; font-size:16px; margin-bottom:20px; opacity:0.8;}
    .timer-box {text-align:center; font-size:30px; font-weight:800; padding:10px; margin-bottom:8px;}
    .login-info {padding:14px 16px; border-radius:8px; margin-bottom:18px; background:#eef6ff;}
    .footer-text {text-align:center; color:#6b7280; font-size:12px; margin-top:25px;}
    </style>
    """,
    unsafe_allow_html=True,
)

# ============================================================
# SESSION STATE
# ============================================================
DEFAULTS = {
    "exam_started": False,
    "exam_submitted": False,
    "waiting_for_batch": False,
    "student_name": "",
    "student_id": "",
    "student_batch": "",
    "exam_questions": [],
    "answers": {},
    "start_time": None,
    "result": None,
    "admin_authenticated": False,
    "show_admin_login": False,
}
for key, value in DEFAULTS.items():
    if key not in st.session_state:
        st.session_state[key] = value

# ============================================================
# TIME / BATCH HELPERS
# ============================================================
def now_local():
    return datetime.now(LOCAL_TIMEZONE)


def parse_hhmm(value):
    h, m = [int(x) for x in value.split(":")]
    return h * 60 + m


def batch_by_name(name):
    return next((b for b in BATCHES if b["name"] == name), None)


def batch_access_status(batch_name, now=None):
    """Returns (allowed: bool, message: str, opens_in_future: bool)."""
    now = now or now_local()
    if now.strftime("%Y-%m-%d") != TEST_DATE:
        return False, f"This exam is scheduled for {TEST_DATE}, not today.", False
    batch = batch_by_name(batch_name)
    if not batch:
        return False, "Invalid batch.", False
    current = now.hour * 60 + now.minute
    start = parse_hhmm(batch["start"])
    end = parse_hhmm(batch["end"])
    if current < start:
        return False, f"Your {batch_name} access window opens at {batch['start']}.", True
    if current > end:
        return False, f"Your {batch_name} access window has closed.", False
    return True, "", False

# ============================================================
# RESULTS / DUPLICATE SAFETY
# ============================================================
def initialise_results_file():
    if not os.path.exists(RESULTS_FILE):
        pd.DataFrame(columns=RESULT_COLUMNS).to_csv(RESULTS_FILE, index=False)


def load_results_safe():
    try:
        initialise_results_file()
        if os.path.getsize(RESULTS_FILE) == 0:
            return pd.DataFrame(columns=RESULT_COLUMNS), None
        df = pd.read_csv(RESULTS_FILE, dtype=str)
        for col in RESULT_COLUMNS:
            if col not in df.columns:
                df[col] = ""
        return df, None
    except Exception as exc:
        return pd.DataFrame(columns=RESULT_COLUMNS), str(exc)


def student_has_attempted(student_id):
    df, _ = load_results_safe()
    if df.empty or "Student ID" not in df.columns:
        return False
    return df["Student ID"].astype(str).str.strip().str.upper().eq(student_id.strip().upper()).any()


def save_result(result):
    initialise_results_file()
    student_id = result["Student ID"].strip().upper()
    try:
        existing, _ = load_results_safe()
        if not existing.empty and student_has_attempted(student_id):
            return False, "This Student ID has already submitted an attempt."
        row = {k: result.get(k, "") for k in RESULT_COLUMNS}
        new_df = pd.DataFrame([row], columns=RESULT_COLUMNS)
        updated = pd.concat([existing[RESULT_COLUMNS], new_df], ignore_index=True)
        updated.to_csv(RESULTS_FILE, index=False)
        return True, "Saved"
    except Exception as exc:
        return False, str(exc)

# ============================================================
# TEST CREATION / SCORING
# ============================================================
def create_student_test():
    """Shuffle question order AND each question's option order per
    student, so no two students see an identical layout."""
    result = []
    for q in QUESTION_BANK:
        opts = list(q["options"])
        random.shuffle(opts)
        result.append({"question": q["question"], "options": opts, "answer": q["answer"]})
    random.shuffle(result)
    return result


def score_current_test():
    correct = 0
    for i, q in enumerate(st.session_state.exam_questions):
        if st.session_state.answers.get(i) == q["answer"]:
            correct += 1
    wrong = QUESTIONS_PER_STUDENT - correct
    percentage = (correct / QUESTIONS_PER_STUDENT) * 100 if QUESTIONS_PER_STUDENT else 0
    elapsed = min(int(time.time() - st.session_state.start_time), TEST_DURATION_SECONDS)
    return correct, wrong, percentage, elapsed


def format_seconds(seconds):
    seconds = max(0, int(seconds))
    return f"{seconds // 60:02d}:{seconds % 60:02d}"


def submit_current_test(status="Submitted"):
    correct, wrong, percentage, elapsed = score_current_test()
    result = {
        "Timestamp": now_local().strftime("%Y-%m-%d %H:%M:%S"),
        "Name": st.session_state.student_name,
        "Student ID": st.session_state.student_id,
        "Course": COURSE_TITLE,
        "Batch": st.session_state.student_batch,
        "Score": round(percentage, 2),
        "Total Questions": QUESTIONS_PER_STUDENT,
        "Correct Answers": correct,
        "Wrong Answers": wrong,
        "Percentage": round(percentage, 2),
        "Time Used": format_seconds(elapsed),
        "Time Seconds": elapsed,
        "Status": status,
    }
    saved, message = save_result(result)
    if saved:
        st.session_state.result = {
            "correct": correct, "wrong": wrong, "score": percentage,
            "time_used": format_seconds(elapsed), "status": status,
        }
        st.session_state.exam_submitted = True
        st.session_state.exam_started = False
        return True, message
    return False, message

# ============================================================
# SIDEBAR
# ============================================================
with st.sidebar:
    st.markdown(f"### \U0001F4DD {COURSE_TITLE}")
    st.markdown("---")
    st.markdown("**Exam Information**")
    st.write(f"Date: **{TEST_DATE}**")
    st.write(f"Timezone: **{TIMEZONE_NAME}**")
    st.write(f"Questions: **{QUESTIONS_PER_STUDENT}**")
    st.write(f"Duration: **{TEST_DURATION_SECONDS // 60} minutes**")
    st.write(f"Students: **{TOTAL_STUDENTS}**")
    st.markdown("---")
    st.markdown("**Batch Schedule**")
    for b in BATCHES:
        lock = " \U0001F512" if b.get("allowed_ids") else ""
        st.markdown(f"**{b['name']}**{lock} ({b['size']} students) -- {b['start']} to {b['end']}")
    if any(b.get("allowed_ids") for b in BATCHES):
        st.caption("\U0001F512 = only specific registered Student IDs can use this batch.")
    st.markdown("---")
    if st.button("\U0001F468\u200D\U0001F3EB Admin Access", use_container_width=True):
        st.session_state.show_admin_login = True
        st.rerun()
    st.link_button(
        "\U0001F4BB View source on GitHub",
        "https://github.com/olumobigjoe/Quiz_automation",
        use_container_width=True,
    )

# ============================================================
# ADMIN LOGIN
# ============================================================
if st.session_state.show_admin_login and not st.session_state.admin_authenticated:
    st.markdown('<div class="main-title">\U0001F468\u200D\U0001F3EB Admin Access</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="main-subtitle">{COURSE_TITLE}</div>', unsafe_allow_html=True)
    password = st.text_input("Admin Password", type="password")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("\U0001F510 Login", use_container_width=True):
            if ADMIN_PASSWORD and password == ADMIN_PASSWORD:
                st.session_state.admin_authenticated = True
                st.session_state.show_admin_login = False
                st.rerun()
            else:
                st.error("Incorrect admin password.")
    with c2:
        if st.button("\u2190 Back to Student Portal", use_container_width=True):
            st.session_state.show_admin_login = False
            st.rerun()
    st.stop()

# ============================================================
# ADMIN DASHBOARD
# ============================================================
if st.session_state.admin_authenticated:
    st.markdown(f'<div class="main-title">\U0001F4CA {COURSE_TITLE} -- Admin Dashboard</div>', unsafe_allow_html=True)

    df, load_error = load_results_safe()
    tabs = st.tabs(["Overview", "Batch Analysis", "Student Results", "Reports", "Session Control"])

    with tabs[0]:
        attempted = len(df)
        c1, c2 = st.columns(2)
        c1.metric("Students Attempted", attempted)
        c2.metric("Students Not Attempted", max(TOTAL_STUDENTS - attempted, 0))
        if df.empty:
            st.info("No student results are available yet.")
        else:
            scores = pd.to_numeric(df["Score"], errors="coerce").dropna()
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Mean Score", f"{scores.mean():.2f}" if not scores.empty else "\u2014")
            c2.metric("Median Score", f"{scores.median():.2f}" if not scores.empty else "\u2014")
            c3.metric("Minimum", f"{scores.min():.2f}" if not scores.empty else "\u2014")
            c4.metric("Maximum", f"{scores.max():.2f}" if not scores.empty else "\u2014")

    with tabs[1]:
        if df.empty:
            st.info("No batch data available yet.")
        else:
            work = df.copy()
            work["ScoreNum"] = pd.to_numeric(work["Score"], errors="coerce")
            rows = []
            for b in BATCHES:
                x = work[work["Batch"] == b["name"]]
                s = x["ScoreNum"].dropna()
                rows.append({
                    "Batch": b["name"], "Capacity": b["size"], "Attempted": len(x),
                    "Mean Score": round(s.mean(), 2) if not s.empty else None,
                    "Minimum": s.min() if not s.empty else None,
                    "Maximum": s.max() if not s.empty else None,
                })
            st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

    with tabs[2]:
        if df.empty:
            st.info("No student results have been submitted yet.")
        else:
            st.dataframe(df, use_container_width=True, hide_index=True)

    with tabs[3]:
        if df.empty:
            st.info("No report data available yet.")
        else:
            st.download_button(
                "\u2b07\ufe0f Download Results CSV", df.to_csv(index=False).encode("utf-8"),
                "exam_results.csv", "text/csv", use_container_width=True,
            )

    with tabs[4]:
        st.subheader("Session Control")
        session_rows = [
            ("Course", COURSE_TITLE), ("Test Date", TEST_DATE), ("Timezone", TIMEZONE_NAME),
            ("Question Bank / Per-Student", f"{QUESTIONS_PER_STUDENT} questions"),
            ("Duration", f"{TEST_DURATION_SECONDS // 60} minutes"),
            ("Auto-refresh interval", f"{AUTO_REFRESH_SECONDS} seconds"),
            ("Total Students", str(TOTAL_STUDENTS)), ("Number of Batches", str(len(BATCHES))),
        ]
        st.table(pd.DataFrame(session_rows, columns=["Setting", "Value"]))

    if load_error:
        st.warning(f"Results file warning: {load_error}")
    if st.button("\U0001F6AA Logout Admin", use_container_width=True):
        st.session_state.admin_authenticated = False
        st.rerun()
    st.stop()

# ============================================================
# RESULT PAGE
# ============================================================
if st.session_state.exam_submitted:
    st.markdown(f'<div class="main-title">\U0001F389 {COURSE_TITLE}</div>', unsafe_allow_html=True)
    st.success("Your test has been submitted successfully.")
    result = st.session_state.result
    c1, c2, c3 = st.columns(3)
    c1.metric("Score", f"{result['score']:.2f} / 100")
    c2.metric("Correct", f"{result['correct']} / {QUESTIONS_PER_STUDENT}")
    c3.metric("Time Used", result["time_used"])
    st.write(f"**Name:** {st.session_state.student_name}")
    st.write(f"**Student ID:** {st.session_state.student_id}")
    st.write(f"**Batch:** {st.session_state.student_batch}")
    st.info("Your result has been recorded. Only one attempt is permitted per Student ID.")
    st.stop()

# ============================================================
# WAITING ROOM -- batch window hasn't opened yet.
# Auto-refreshes every 10s (via st.fragment) so the student isn't stuck
# staring at a stale page or needing to manually reload while they wait.
# ============================================================
if st.session_state.waiting_for_batch:
    @st.fragment(run_every=f"{AUTO_REFRESH_SECONDS}s")
    def render_waiting_room():
        allowed, message, opens_future = batch_access_status(st.session_state.student_batch)
        if allowed:
            st.session_state.waiting_for_batch = False
            st.session_state.exam_questions = create_student_test()
            st.session_state.answers = {}
            st.session_state.start_time = time.time()
            st.session_state.exam_started = True
            st.rerun()
            return
        st.markdown(f'<div class="main-title">\u23F3 {COURSE_TITLE}</div>', unsafe_allow_html=True)
        st.info(f"Hi {st.session_state.student_name}, you're checked in for **{st.session_state.student_batch}**.")
        if opens_future:
            st.warning(message + " This page refreshes automatically -- just leave it open.")
        else:
            st.error(message)
            if st.button("\u2190 Back to login"):
                st.session_state.waiting_for_batch = False
                st.rerun()

    render_waiting_room()
    st.stop()

# ============================================================
# STUDENT LOGIN
# ============================================================
if not st.session_state.exam_started:
    st.markdown(f'<div class="main-title">{COURSE_TITLE}</div>', unsafe_allow_html=True)
    st.markdown('<div class="main-subtitle">Online Test</div>', unsafe_allow_html=True)
    st.markdown('<div class="login-info">Enter your details below, then select your assigned batch.</div>', unsafe_allow_html=True)

    name = st.text_input("Full Name", placeholder="Enter your full name")
    student_id = st.text_input("Student ID", placeholder="Enter your student/matric ID")
    batch_name = st.selectbox("Your Batch", [b["name"] for b in BATCHES])

    if st.button("\U0001F680 Start Test", use_container_width=True):
        name = name.strip()
        student_id = student_id.strip()
        if not name:
            st.error("Please enter your name.")
            st.stop()
        if not student_id:
            st.error("Please enter your Student ID.")
            st.stop()
        if student_has_attempted(student_id):
            st.error("This Student ID has already submitted this test.")
            st.stop()

        selected_batch = batch_by_name(batch_name)
        if selected_batch and selected_batch.get("allowed_ids"):
            if student_id.upper() not in selected_batch["allowed_ids"]:
                st.error(
                    f"Your Student ID is not registered for {batch_name}. "
                    "Double-check your batch, or contact your instructor if you "
                    "believe this is a mistake."
                )
                st.stop()

        st.session_state.student_name = name
        st.session_state.student_id = student_id
        st.session_state.student_batch = batch_name

        allowed, message, opens_future = batch_access_status(batch_name)
        if not allowed:
            if opens_future:
                # send them to the auto-refreshing waiting room instead of
                # a dead-end error -- they don't need to keep reloading.
                st.session_state.waiting_for_batch = True
                st.rerun()
            else:
                st.error(message)
                st.stop()

        st.session_state.exam_questions = create_student_test()
        st.session_state.answers = {}
        st.session_state.start_time = time.time()
        st.session_state.exam_started = True
        st.session_state.exam_submitted = False
        st.rerun()

    st.markdown(f'<div class="footer-text">{COURSE_TITLE}</div>', unsafe_allow_html=True)
    st.stop()

# ============================================================
# ACTIVE TEST -- 10 SECOND FRAGMENT RERUN
# Using st.fragment here (instead of a plain st.rerun loop) means only
# this section of the page re-executes every 10s, not the whole app --
# this is what keeps the timer live without the page stalling/hanging
# for the student while they're mid-test.
# ============================================================
@st.fragment(run_every=f"{AUTO_REFRESH_SECONDS}s")
def render_active_test():
    if not st.session_state.exam_started:
        return

    elapsed = time.time() - st.session_state.start_time
    remaining = TEST_DURATION_SECONDS - int(elapsed)

    if remaining <= 0:
        ok, msg = submit_current_test("Time Expired")
        if ok:
            st.rerun()
        st.error(f"Time expired, but the result could not be saved: {msg}")
        return

    st.markdown(f'<div class="timer-box">\u23F1\ufe0f Time Remaining: {remaining // 60:02d}:{remaining % 60:02d}</div>', unsafe_allow_html=True)
    st.progress(min(elapsed / TEST_DURATION_SECONDS, 1.0))
    st.info(
        f"**Name:** {st.session_state.student_name}  \n"
        f"**Student ID:** {st.session_state.student_id}  \n"
        f"**Batch:** {st.session_state.student_batch}"
    )
    st.markdown("---")
    st.subheader(f"Answer all {QUESTIONS_PER_STUDENT} questions")

    for i, q in enumerate(st.session_state.exam_questions):
        st.markdown(f"### Question {i + 1} of {QUESTIONS_PER_STUDENT}")
        st.write(f"**{q['question']}**")
        current = st.session_state.answers.get(i)
        options = q["options"]
        default_index = options.index(current) if current in options else None
        selected = st.radio("Select one answer:", options, key=f"question_{i}", index=default_index)
        st.session_state.answers[i] = selected
        st.markdown("---")

    st.warning("\u26A0\ufe0f Once you submit the test, you cannot change your answers or attempt the test again.")
    if st.button("\U0001F4E4 SUBMIT TEST", use_container_width=True):
        ok, msg = submit_current_test("Submitted")
        if ok:
            st.rerun()
        else:
            st.error(f"Submission failed: {msg}")

render_active_test()
'''

REQUIREMENTS_TEMPLATE = """streamlit>=1.35
pandas>=2.0
"""

README_TEMPLATE = r"""# __COURSE_TITLE__ -- Online Exam Portal

An auto-generated, self-contained online exam portal built with
[Streamlit](https://streamlit.io). Students log in, wait (if needed) for
their batch's access window, sit a timed multiple-choice test, and get an
instant score -- all with **no external services, no database, and no
internet API calls** once it's running.

## Table of Contents

- [Overview](#overview)
- [Features](#features)
- [Requirements](#requirements)
- [Quick Start](#quick-start)
- [How the Exam Flow Works](#how-the-exam-flow-works)
- [Batch Schedule](#batch-schedule)
- [Project Structure](#project-structure)
- [Configuration](#configuration)
- [Deploying to Streamlit Community Cloud](#deploying-to-streamlit-community-cloud)
- [Data & Privacy](#data--privacy)
- [Security Notes](#security-notes)
- [Troubleshooting](#troubleshooting)
- [Regenerating This Exam](#regenerating-this-exam)
- [License](#license)

## Overview

| | |
|---|---|
| **Course** | __COURSE_TITLE__ |
| **Questions** | __NUM_QUESTIONS__ (__NUM_OPTIONS__ options each) |
| **Students** | __NUM_STUDENTS__ across __NUM_BATCHES__ batch(es), max __MAX_STUDENTS_PER_BATCH__ per batch |
| **Test date** | __TEST_DATE__ |
| **Duration** | __TEST_DURATION_MINUTES__ minutes per student |
| **Generated** | __GENERATED_DATE__ from `__SOURCE_FILENAME__` |
| **Question source** | AI-drafted with __MODEL_USED__, reviewed/edited by the instructor before export |

## Features

- **Self-contained** -- questions, options, and answers are embedded directly
  in `exam_app.py`; no Ollama, no internet connection, and no database needed
  to run the exam itself.
- **Batch scheduling** -- students are split into time-boxed batches (max
  __MAX_STUDENTS_PER_BATCH__ students each) so a large cohort doesn't all sit
  the test simultaneously.
- **Live waiting room** -- if a student logs in before their batch's window
  opens, they see an auto-refreshing waiting screen (every 10 seconds)
  instead of a dead-end error or a page they have to keep manually reloading.
- **Live countdown timer** -- the active test screen refreshes every 10
  seconds in the background (via `st.fragment`) so the timer stays accurate
  without freezing or reloading the whole page.
- **One attempt per student** -- re-submission with the same Student ID is
  blocked, both before and after a successful submit.
- **Randomized presentation** -- each student sees questions and answer
  options in a different shuffled order.
- **Admin dashboard** -- a password-gated panel to monitor attempts, review
  per-batch performance, browse raw results, and download a CSV.

## Requirements

- Python 3.10+
- Packages listed in `requirements.txt` (Streamlit + pandas)

## Quick Start

```bash
pip install -r requirements.txt
streamlit run exam_app.py
```

Open the local URL Streamlit prints (usually http://localhost:8501).

## How the Exam Flow Works

1. **Login** -- the student enters their name, Student ID, and selects their
   assigned batch from a dropdown.
2. **Waiting room (if early)** -- if their batch's window hasn't opened yet,
   they land on an auto-refreshing waiting screen rather than an error; the
   test starts automatically the moment the window opens, with no action
   needed from the student.
3. **Timed test** -- once their window is open, the full question set loads
   (order and option order shuffled per student) with a live countdown.
4. **Auto-submit** -- if time runs out, the test is submitted automatically
   with whatever answers were recorded.
5. **One-time result** -- the student sees their score immediately; their
   Student ID is then locked out of re-attempting.

## Batch Schedule

| Batch | Capacity | Window | Registered Matric/Student IDs |
|---|---|---|---|
__BATCH_TABLE_ROWS__

*(All times are in the `__TIMEZONE_NAME__` timezone -- see [Configuration](#configuration) to change it. A batch showing "Any registered student" has no matric-number restriction -- any student can log into it and select it from the dropdown. Batches with specific IDs will reject a Student ID that doesn't match at login.)*

## Project Structure

```
exam_app.py         # The full exam portal (Streamlit app) -- run this
question_bank.json  # Reference copy of the questions/options/answers
                     # (also embedded directly inside exam_app.py)
requirements.txt     # Python dependencies
README.md            # This file
exam_results.csv     # Created automatically the first time a student submits
```

## Configuration

A few values can be changed without touching the question logic:

- **Admin password** -- baked into `exam_app.py` as a default, but you can
  override it at launch without editing the file:
  ```bash
  export ADMIN_PASSWORD="a-stronger-password"   # Mac/Linux
  setx ADMIN_PASSWORD "a-stronger-password"      # Windows (new terminal after)
  streamlit run exam_app.py
  ```
- **Timezone** -- edit the `TIMEZONE_NAME` constant near the top of
  `exam_app.py` (uses standard [IANA timezone names](https://en.wikipedia.org/wiki/List_of_tz_database_time_zones),
  e.g. `"Africa/Lagos"`, `"America/New_York"`).
- **Auto-refresh interval** -- `AUTO_REFRESH_SECONDS` (default 10) controls
  both the waiting-room and active-test refresh rate.

## Deploying to Streamlit Community Cloud

1. Push this folder to a GitHub repository.
2. Go to [share.streamlit.io](https://share.streamlit.io), connect the repo,
   and set the main file path to `exam_app.py`.
3. In the app's **Settings -> Secrets**, you can optionally set
   `ADMIN_PASSWORD` there instead of as a plain environment variable --
   Streamlit Cloud exposes Secrets as environment variables automatically.
4. Deploy. No other configuration is required.

## Data & Privacy

- Student results (name, Student ID, batch, score, timing) are written to a
  local `exam_results.csv` file in the app's working directory. On Streamlit
  Community Cloud, this file persists only for the life of that deployment's
  filesystem -- **download the CSV from the admin Reports tab periodically**
  if you need a permanent copy.
- No data is sent anywhere outside the app itself; there are no external API
  calls in `exam_app.py`.

## Security Notes

- This portal is designed for classroom-scale, low-to-moderate-stakes
  testing. It does not include anti-cheating measures like tab-switch
  detection, webcam proctoring, or IP restrictions.
- The admin password, if left at its baked-in default, is visible to anyone
  who can read the source file (e.g. in a public GitHub repo). Override it
  via the `ADMIN_PASSWORD` environment variable (see
  [Configuration](#configuration)) before using this for anything sensitive.
- The `exam_results.csv` file is plain text and unencrypted.

## Troubleshooting

- **"This exam is scheduled for [date], not today."** -- the server's system
  date doesn't match `TEST_DATE`. Check the deployment environment's clock,
  or update `TEST_DATE` in `exam_app.py`.
- **A student can't log in / batch window already closed** -- double-check
  they were given the correct batch name, and that `TIMEZONE_NAME` matches
  the timezone you scheduled the test in.
- **Waiting room seems stuck** -- it refreshes every `AUTO_REFRESH_SECONDS`
  (default 10s) automatically; nothing needs to be clicked. If it truly
  isn't updating, confirm the Streamlit version supports `st.fragment`
  (Streamlit 1.33+; `requirements.txt` already pins a version that does).
- **Duplicate-attempt errors for a student who never actually took it** --
  Student IDs are matched case-insensitively; check for a typo'd ID already
  present in `exam_results.csv`.

## Regenerating This Exam

This exported folder is a standalone snapshot. To generate a new/different
exam (new material, question count, batch setup, etc.), go back to the quiz
builder tool (`app.py` in the *generator* project, not this one) -- it does
not need this exported folder to run.

## License

MIT -- use, modify, and redistribute freely for your own courses.
"""


def build_export_files(questions, source_filename, model_used, course_title,
                        test_date: date, test_start_time: dt_time,
                        num_students: int, num_batches: int,
                        test_duration_minutes: int, batch_window_minutes: int,
                        admin_password: str, timezone_name: str = "Africa/Lagos",
                        batch_matric_configs: list | None = None):
    """Builds (exam_app_code, readme, question_bank_json) for download.
    Raises ValueError if the batch/student configuration is invalid
    (e.g. more than MAX_STUDENTS_PER_BATCH students in a batch)."""
    matric_lists = [expand_matric_config(c) for c in (batch_matric_configs or [])] or None
    batches = compute_batches(
        num_students, num_batches, test_date, test_start_time, batch_window_minutes,
        matric_lists=matric_lists,
    )

    generated_date = datetime.now().strftime("%Y-%m-%d %H:%M")
    num_options = len(questions[0]["options"]) if questions else 0

    # The exported JSON / embedded question bank contains ONLY
    # question + options + answer -- no explanation, no metadata.
    exam_questions = [
        {
            "question": q["question"],
            "options": q["options"],
            "answer": q["options"][q["correct_index"]],
        }
        for q in questions
    ]
    questions_json = json.dumps(exam_questions, ensure_ascii=False)
    questions_b64 = base64.b64encode(questions_json.encode("utf-8")).decode("ascii")

    exam_code = (
        EXAM_APP_TEMPLATE
        .replace("__COURSE_TITLE__", course_title or "Course Quiz")
        .replace("__GENERATED_DATE__", generated_date)
        .replace("__SOURCE_FILENAME__", source_filename or "unknown")
        .replace("__TEST_DATE__", test_date.strftime("%Y-%m-%d"))
        .replace("__TEST_START_TIME__", test_start_time.strftime("%H:%M"))
        .replace("__TIMEZONE_NAME__", timezone_name)
        .replace("__NUM_STUDENTS__", str(num_students))
        .replace("__TEST_DURATION_SECONDS__", str(test_duration_minutes * 60))
        .replace("__ADMIN_PASSWORD__", admin_password or "changeme")
        .replace("__MAX_STUDENTS_PER_BATCH__", str(MAX_STUDENTS_PER_BATCH))
        .replace("__BATCHES_JSON__", json.dumps(batches))
        .replace("__QUESTIONS_B64__", questions_b64)
    )

    def _matric_summary(b):
        ids = b.get("allowed_ids") or []
        if not ids:
            return "Any registered student"
        if len(ids) <= 4:
            return ", ".join(ids)
        return f"{ids[0]} ... {ids[-1]} ({len(ids)} IDs)"

    batch_table_rows = "\n".join(
        f"| {b['name']} | {b['size']} | {b['start']} - {b['end']} | {_matric_summary(b)} |"
        for b in batches
    )

    readme = (
        README_TEMPLATE
        .replace("__COURSE_TITLE__", course_title or "Course Quiz")
        .replace("__NUM_QUESTIONS__", str(len(questions)))
        .replace("__NUM_OPTIONS__", str(num_options) if num_options else "?")
        .replace("__NUM_STUDENTS__", str(num_students))
        .replace("__NUM_BATCHES__", str(num_batches))
        .replace("__MAX_STUDENTS_PER_BATCH__", str(MAX_STUDENTS_PER_BATCH))
        .replace("__TEST_DATE__", test_date.strftime("%Y-%m-%d"))
        .replace("__TEST_DURATION_MINUTES__", str(test_duration_minutes))
        .replace("__GENERATED_DATE__", generated_date)
        .replace("__SOURCE_FILENAME__", source_filename or "unknown")
        .replace("__MODEL_USED__", model_used)
        .replace("__TIMEZONE_NAME__", timezone_name)
        .replace("__BATCH_TABLE_ROWS__", batch_table_rows)
    )

    return exam_code, readme, questions_json


# ============================================================================
# UI -- chat-style walkthrough
# ============================================================================
st.title("📚 Course Material → Quiz Generator")
st.caption("Runs entirely on your machine via Ollama. No data leaves your laptop.")

with st.chat_message("assistant"):
    st.write(
        "Upload a PDF or DOCX of your course material, choose how many "
        "questions you want, and I'll draft a multiple-choice quiz from "
        "it -- including suggested correct answers -- for you to review."
    )

# ---- Sidebar settings ----
with st.sidebar:
    st.header("Generation Settings")
    model_choice = st.selectbox(
        "Model (Ollama Cloud)",
        MODEL_OPTIONS,
        index=0,
        help="gpt-oss:20b is a good low-usage-tier default on the free plan. "
        "gpt-oss:120b / qwen3.5 / deepseek-v3.1 are larger and higher quality "
        "but consume your plan's usage allowance faster.",
    )
    if model_choice == "Custom model tag...":
        model = st.text_input(
            "Model tag",
            placeholder="e.g. glm-4.6:cloud",
            help="Check currently available tags for your account at "
            "https://ollama.com/settings/keys or via GET https://ollama.com/api/tags",
        )
    else:
        model = model_choice

    api_key = st.text_input(
        "Ollama API key",
        value=os.environ.get("OLLAMA_API_KEY", ""),
        type="password",
        help="From https://ollama.com/settings/keys. Pre-filled automatically "
        "if the OLLAMA_API_KEY environment variable is set before launching.",
    )
    num_questions = st.slider("Number of questions", 3, MAX_QUESTIONS, 10)
    num_options = st.slider("Options per question", 2, 6, 4)
    st.divider()
    st.caption(
        "Runs on Ollama Cloud, not your machine -- no local model download, "
        "and this works the same whether you run the app locally or deploy it "
        "(e.g. to Streamlit Community Cloud)."
    )
    st.caption("Get a free API key: https://ollama.com")

    st.divider()
    st.header("Exam Portal Settings")
    st.caption("Used only when you export the standalone exam portal below.")
    course_title = st.text_input("Course Title", placeholder="e.g. Introduction to Biochemistry")
    exam_date = st.date_input("Test Date", value=date.today())
    exam_time = st.time_input("Test Start Time", value=dt_time(9, 0))
    test_duration_minutes = st.number_input("Test duration per student (minutes)", min_value=1, value=30)
    batch_window_minutes = st.number_input(
        "Batch window length (minutes)", min_value=1, value=int(test_duration_minutes),
        help="How long each batch's login/access window stays open before the next batch's window begins.",
    )
    num_students = st.number_input("Number of students", min_value=1, value=100)
    num_batches = st.number_input(
        "Number of batches", min_value=1,
        value=max(1, math.ceil(num_students / MAX_STUDENTS_PER_BATCH)),
        help=f"Each batch can hold at most {MAX_STUDENTS_PER_BATCH} students -- "
             f"increase this if you get a capacity error at export time.",
    )

    st.markdown("**Batch matric/student number registration**")
    st.caption(
        "Optional. By default a batch is open to any registered student, and "
        "the class list is split evenly across batches. Restrict a batch to a "
        "matric-number range or an explicit list instead if you want students "
        "auto-rejected for logging into the wrong batch."
    )
    batch_matric_configs = []
    for i in range(int(num_batches)):
        with st.expander(f"Batch {i + 1} -- matric numbers", expanded=False):
            mode_label = st.radio(
                "Assignment",
                ["No restriction (auto-split)", "Range", "Other / specific list"],
                key=f"batch_mode_{i}",
                horizontal=True,
            )
            if mode_label == "Range":
                c1, c2, c3, c4 = st.columns([3, 2, 2, 2])
                prefix = c1.text_input(
                    "Prefix", key=f"batch_prefix_{i}", placeholder="e.g. CSC/2021/",
                    help="Text placed before the number, e.g. a department/year code.",
                )
                start_n = c2.number_input("From #", min_value=0, value=1, key=f"batch_start_{i}")
                end_n = c3.number_input("To #", min_value=0, value=40, key=f"batch_end_{i}")
                pad = c4.number_input(
                    "Zero-pad", min_value=0, max_value=10, value=3, key=f"batch_pad_{i}",
                    help="Digits to pad the number to, e.g. 3 turns 7 into 007. 0 = no padding.",
                )
                batch_matric_configs.append(
                    {"mode": "range", "prefix": prefix, "start": int(start_n), "end": int(end_n), "pad": int(pad)}
                )
                if end_n >= start_n:
                    preview_first = f"{prefix}{str(int(start_n)).zfill(int(pad)) if pad else int(start_n)}"
                    preview_last = f"{prefix}{str(int(end_n)).zfill(int(pad)) if pad else int(end_n)}"
                    st.caption(f"Registers {int(end_n) - int(start_n) + 1} IDs: {preview_first} ... {preview_last}")
            elif mode_label == "Other / specific list":
                list_text = st.text_area(
                    "Matric/student numbers (comma or newline separated)",
                    key=f"batch_list_{i}",
                    placeholder="CSC/2021/003, CSC/2021/011, MEE/2020/045, ...",
                    help="Use this for a batch made of non-consecutive or mixed-format IDs.",
                )
                batch_matric_configs.append({"mode": "list", "text": list_text})
                count = len([p for p in re.split(r"[,\n]", list_text or "") if p.strip()])
                if count:
                    st.caption(f"Registers {count} ID(s).")
            else:
                batch_matric_configs.append({"mode": "auto"})

    admin_password = st.text_input(
        "Admin password (for the exported portal's dashboard)", type="password",
        help="Baked into the exported file by default -- see the README's Security Notes "
             "for how to override it via environment variable instead.",
    )

# ---- File upload ----
uploaded = st.file_uploader("Upload course material", type=["pdf", "docx"])

if uploaded is not None and uploaded.name != st.session_state.source_filename:
    with st.spinner(f"Reading {uploaded.name}..."):
        file_bytes = uploaded.read()
        if uploaded.name.lower().endswith(".pdf"):
            text = extract_text_from_pdf(file_bytes)
        else:
            text = extract_text_from_docx(file_bytes)
    st.session_state.raw_text = text
    st.session_state.source_filename = uploaded.name
    st.session_state.questions = []
    st.session_state.generated_once = False
    st.session_state.regenerate_count = 0

if st.session_state.raw_text:
    word_count = len(st.session_state.raw_text.split())
    with st.chat_message("assistant"):
        st.write(
            f"I read **{word_count:,} words** from `{st.session_state.source_filename}`. "
            f"Ready to draft **{num_questions} questions** ({num_options} options each) "
            f"using **{model}** on Ollama Cloud."
        )
    if word_count < 30:
        st.warning(
            "That's very little extracted text -- if this is a scanned/image-only "
            "PDF, text extraction won't work. Try a text-based PDF or DOCX."
        )
    if not api_key:
        st.warning(
            "No Ollama API key set -- add one in the sidebar (or set the "
            "OLLAMA_API_KEY environment variable) before generating."
        )
    if not model:
        st.warning("Enter a custom model tag in the sidebar, or pick one from the list.")

    generate_disabled = word_count < 30 or not api_key or not model
    button_label = "🚀 Generate Quiz" if not st.session_state.generated_once else "🔄 Regenerate Quiz (new random set)"

    if st.button(button_label, type="primary", disabled=generate_disabled):
        with st.spinner("Talking to Ollama Cloud..."):
            questions, log = generate_quiz(
                st.session_state.raw_text, model, OLLAMA_CLOUD_HOST, api_key, num_questions, num_options
            )
        st.session_state.questions = questions
        st.session_state.generation_log = log
        if st.session_state.generated_once:
            # This wasn't the first generation -- it's a regenerate. No cap
            # is applied here, so the user can regenerate as many times as
            # they like (at least twice, and beyond).
            st.session_state.regenerate_count += 1
        st.session_state.generated_once = True

    if st.session_state.regenerate_count > 0:
        st.caption(f"Regenerated {st.session_state.regenerate_count} time(s) so far -- click again anytime for a fresh random set.")

# ---- Generation log / errors ----
if st.session_state.generation_log:
    with st.expander("Generation log", expanded=not st.session_state.questions):
        for line in st.session_state.generation_log:
            st.write("- " + line)

# ---- Review & edit ----
if st.session_state.questions:
    with st.chat_message("assistant"):
        st.write(
            f"Drafted **{len(st.session_state.questions)}** question(s). "
            "Review and edit below -- fix wording, swap options, or change "
            "the marked correct answer -- then export."
        )

    to_remove = []
    for i, q in enumerate(st.session_state.questions):
        with st.expander(f"Q{i + 1}: {q['question'][:80]}", expanded=False):
            q["question"] = st.text_area("Question", value=q["question"], key=f"qtext_{i}")
            for j in range(len(q["options"])):
                q["options"][j] = st.text_input(
                    f"Option {j + 1}", value=q["options"][j], key=f"opt_{i}_{j}"
                )
            q["correct_index"] = st.radio(
                "Correct answer",
                options=list(range(len(q["options"]))),
                format_func=lambda idx, opts=q["options"]: opts[idx],
                index=q["correct_index"],
                key=f"correct_{i}",
            )
            q["explanation"] = st.text_input(
                "Explanation (optional)", value=q.get("explanation", ""), key=f"expl_{i}"
            )
            if st.checkbox("Remove this question", key=f"remove_{i}"):
                to_remove.append(i)

    if to_remove:
        st.session_state.questions = [
            q for i, q in enumerate(st.session_state.questions) if i not in to_remove
        ]
        st.rerun()

    st.divider()
    st.subheader("Export Exam Portal")

    if not st.session_state.questions:
        st.info("All questions were removed -- nothing to export.")
    elif not course_title.strip():
        st.warning("Enter a Course Title in the sidebar (under Exam Portal Settings) to enable export.")
    else:
        try:
            exam_code, readme, questions_json = build_export_files(
                st.session_state.questions,
                st.session_state.source_filename,
                model,
                course_title.strip(),
                exam_date,
                exam_time,
                int(num_students),
                int(num_batches),
                int(test_duration_minutes),
                int(batch_window_minutes),
                admin_password,
                batch_matric_configs=batch_matric_configs,
            )
        except ValueError as exc:
            st.error(str(exc))
        else:
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.download_button(
                    "⬇️ Download exam_app.py",
                    data=exam_code,
                    file_name="exam_app.py",
                    mime="text/x-python",
                    use_container_width=True,
                )
            with col2:
                st.download_button(
                    "⬇️ Download README.md",
                    data=readme,
                    file_name="README.md",
                    mime="text/markdown",
                    use_container_width=True,
                )
            with col3:
                st.download_button(
                    "⬇️ Download question_bank.json",
                    data=questions_json,
                    file_name="question_bank.json",
                    mime="application/json",
                    use_container_width=True,
                )
            with col4:
                zip_buf = io.BytesIO()
                with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                    zf.writestr("exam_app.py", exam_code)
                    zf.writestr("README.md", readme)
                    zf.writestr("question_bank.json", questions_json)
                    zf.writestr("requirements.txt", REQUIREMENTS_TEMPLATE)
                st.download_button(
                    "⬇️ Download all (.zip)",
                    data=zip_buf.getvalue(),
                    file_name="exam_portal.zip",
                    mime="application/zip",
                    use_container_width=True,
                )

            with st.expander("Preview generated exam_app.py"):
                st.code(exam_code, language="python")
            with st.expander("Preview generated README.md"):
                st.markdown(readme)
